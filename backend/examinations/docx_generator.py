"""
DOCX generator for exam papers.
Generates formatted question papers with school metadata and snapshot-backed questions.

Legacy (empty-structure) papers render through the original flat, unstructured
layout unchanged. Structured papers (ExamPaper.structure non-empty) render through
the shared layout plan in paper_export_layout.py in the classic school-paper format
(header block, section headings with marks, per-type question rendering, answer
lines, deterministic matching-table shuffle) -- kept consistent with the PDF export.
"""

import io
import logging
import re
from datetime import datetime
from urllib.parse import urlparse

from django.utils.html import strip_tags
import requests

from .paper_export_layout import build_export_layout

logger = logging.getLogger(__name__)


def _html_to_text(value):
    if not value:
        return ''
    text = re.sub(r'<br\s*/?>', '\n', str(value), flags=re.IGNORECASE)
    text = re.sub(r'</p\s*>', '\n\n', text, flags=re.IGNORECASE)
    text = strip_tags(text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


class ExamPaperDOCXGenerator:
    """Generate .docx exam papers with snapshot-first question rendering."""

    def __init__(self, exam_paper):
        self.exam_paper = exam_paper
        self.school = exam_paper.school

    def _exam_name(self):
        if self.exam_paper.exam:
            return self.exam_paper.exam.name
        if self.exam_paper.exam_subject and self.exam_paper.exam_subject.exam:
            return self.exam_paper.exam_subject.exam.name
        return None

    def _append_school_logo(self, document, width_inches):
        logo_url = getattr(self.school, 'logo_url', None)
        if not logo_url:
            return

        parsed = urlparse(str(logo_url))
        if not parsed.scheme or not parsed.netloc:
            return

        try:
            response = requests.get(logo_url, timeout=8)
            response.raise_for_status()
            logo_stream = io.BytesIO(response.content)

            from docx.enum.text import WD_ALIGN_PARAGRAPH
            logo_paragraph = document.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_stream, width=width_inches)
        except Exception as exc:
            logger.warning('Could not attach school logo to DOCX for paper %s: %s', self.exam_paper.id, exc)

    def generate(self):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
        except ImportError:
            logger.error('python-docx not installed, cannot generate DOCX')
            raise ImportError('python-docx is required for DOCX generation')

        document = Document()

        self._append_school_logo(document, Inches(1.15))

        layout = build_export_layout(self.exam_paper)
        if layout is None:
            self._render_legacy(document, WD_ALIGN_PARAGRAPH)
        else:
            self._render_structured(document, layout, WD_ALIGN_PARAGRAPH, Inches)

        output = io.BytesIO()
        document.save(output)
        logger.info('Generated DOCX for ExamPaper %s', self.exam_paper.id)
        return output.getvalue()

    def _render_legacy(self, document, align):
        """Unchanged from the original flat-list rendering -- legacy (empty-structure)
        papers must keep exporting exactly as before."""
        school_heading = document.add_heading(self.school.name, level=1)
        school_heading.alignment = align.CENTER

        exam_name = self._exam_name()
        if exam_name:
            exam_heading = document.add_heading(f'Exam: {exam_name}', level=3)
            exam_heading.alignment = align.CENTER

        paper_heading = document.add_heading(self.exam_paper.paper_title, level=2)
        paper_heading.alignment = align.CENTER

        meta_lines = [
            f"Class: {self.exam_paper.class_obj.name}",
            f"Subject: {self.exam_paper.subject.name}",
            f"Total Marks: {self.exam_paper.total_marks}",
            f"Duration: {self.exam_paper.duration_minutes} minutes",
        ]
        if exam_name:
            meta_lines.append(f"Exam: {exam_name}")
        document.add_paragraph(' | '.join(meta_lines))

        if self.exam_paper.instructions:
            document.add_paragraph('Instructions:')
            instructions = _html_to_text(self.exam_paper.instructions)
            for line in [entry.strip() for entry in instructions.splitlines() if entry.strip()]:
                document.add_paragraph(line, style='List Bullet')

        paper_questions = self.exam_paper.paper_questions.select_related('question').order_by('question_order')
        for paper_question in paper_questions:
            question = paper_question.get_question_data()
            marks = paper_question.get_marks()

            heading_text = f"Q{paper_question.question_order}. ({marks} marks)"
            document.add_paragraph(heading_text)

            question_text = _html_to_text(question.get('question_text'))
            document.add_paragraph(question_text or '-')

            if question.get('question_type') == 'MCQ':
                for option_key in ('A', 'B', 'C', 'D'):
                    option_value = question.get(f'option_{option_key.lower()}')
                    if option_value:
                        document.add_paragraph(f"{option_key}. {_html_to_text(option_value)}")
            else:
                document.add_paragraph('')

        footer = document.add_paragraph(
            f"Generated on {datetime.now().strftime('%d %B %Y')} | {self.school.name}"
        )
        footer.alignment = align.CENTER

    def _render_structured(self, document, layout, align, inches):
        """Classic school-paper format for structured papers (non-empty ExamPaper.structure)."""
        from docx.enum.text import WD_TAB_ALIGNMENT

        header = layout['header']

        school_heading = document.add_heading(header['school_name'], level=1)
        school_heading.alignment = align.CENTER

        if header['exam_name']:
            exam_heading = document.add_heading(f"Exam: {header['exam_name']}", level=3)
            exam_heading.alignment = align.CENTER

        paper_heading = document.add_heading(header['paper_title'], level=2)
        paper_heading.alignment = align.CENTER

        subject_class_p = document.add_paragraph(
            f"Paper: {header['subject_name']}    Class: {header['class_name']}"
        )
        subject_class_p.alignment = align.CENTER

        document.add_paragraph(f"Name: {'_' * 40}")
        document.add_paragraph(f"Roll No: {'_' * 20}     Date: {'_' * 15}")
        document.add_paragraph(
            f"Total Marks: {header['total_marks']}     Time: {header['duration_minutes']} minutes"
        )

        if header['instructions']:
            document.add_paragraph('Instructions:')
            instructions_text = _html_to_text(header['instructions'])
            for line in [entry.strip() for entry in instructions_text.splitlines() if entry.strip()]:
                document.add_paragraph(line, style='List Bullet')

        for block in layout['blocks']:
            if block['type'] == 'divider':
                self._render_divider_heading(document, block, align)
                continue
            if block['type'] == 'section':
                self._render_section_heading(document, block, inches, WD_TAB_ALIGNMENT)
            for item in block['items']:
                self._render_question_item(document, item)

        footer = document.add_paragraph(
            f"Generated on {datetime.now().strftime('%d %B %Y')} | {header['school_name']}"
        )
        footer.alignment = align.CENTER

    def _render_divider_heading(self, document, block, align):
        """A plain print-layout separator (e.g. 'Section A') -- no marks, no questions."""
        heading = document.add_heading(block['title'], level=2)
        heading.alignment = align.CENTER

    def _render_section_heading(self, document, block, inches, tab_alignment):
        heading_text = block['title']
        if block['instruction']:
            heading_text = f"{heading_text}. {block['instruction']}"

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.tab_stops.add_tab_stop(inches(6.5), tab_alignment.RIGHT)
        run = paragraph.add_run(heading_text)
        run.bold = True
        paragraph.add_run(f"\t({block['section_marks']})")

    def _render_question_item(self, document, item):
        document.add_paragraph(f"Q{item['number']}. ({item['marks']} marks)")

        question_text = _html_to_text(item['question_text'])
        document.add_paragraph(question_text or '-')

        rendered_extra = False

        if item['question_type'] == 'MCQ' and item['options']:
            for option_key in ('A', 'B', 'C', 'D'):
                option_value = item['options'].get(option_key)
                if option_value:
                    document.add_paragraph(f"{option_key}. {_html_to_text(option_value)}")
            rendered_extra = True

        elif item['question_type'] == 'FILL_BLANK' and item['fill_blank_items']:
            for blank_line in item['fill_blank_items']:
                document.add_paragraph(blank_line)
            rendered_extra = True

        elif item['question_type'] == 'MATCHING' and item['matching_pairs']:
            self._render_matching_table(document, item['matching_pairs'])
            rendered_extra = True

        elif item['question_type'] == 'TRUE_FALSE':
            document.add_paragraph('True / False')
            rendered_extra = True

        if item['answer_lines']:
            for _ in range(item['answer_lines']):
                document.add_paragraph('_' * 60)
            rendered_extra = True

        if not rendered_extra:
            document.add_paragraph('')

    def _render_matching_table(self, document, pairs):
        table = document.add_table(rows=len(pairs) + 1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Column A'
        header_cells[1].text = 'Column B'
        for row_index, pair in enumerate(pairs, start=1):
            cells = table.rows[row_index].cells
            cells[0].text = pair['left']
            cells[1].text = pair['right']
